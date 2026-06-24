"""Generate synthetic .docx fixtures (resume + cover letter) for the
Jason Statham test applicant, sourced from jason_statham.json.

Run:  python test_assets/generate_docx.py
Produces: jason_statham_resume.docx, jason_statham_cover_letter.docx
These are SYNTHETIC fixtures tailored to the FlashTec Senior .NET Engineer
posting (FT-2024-8842). Not real personal data.
"""
import json
import os
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
DATA = json.load(open(os.path.join(HERE, "jason_statham.json"), encoding="utf-8"))

ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def _name_block(doc, subtitle):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(DATA["fullName"])
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle)
    sr.font.size = Pt(11)
    sr.font.color.rgb = ACCENT

    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact = " | ".join([
        DATA["location"],
        DATA["phone"],
        DATA["email"],
        DATA["linkedinUrl"].replace("https://www.", "").replace("https://", ""),
        DATA["githubUrl"].replace("https://", ""),
    ])
    cr = c.add_run(contact)
    cr.font.size = Pt(9)


def _section(doc, title):
    p = doc.add_paragraph()
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    # underline rule
    pr = doc.add_paragraph()
    pr.paragraph_format.space_after = Pt(4)
    rr = pr.add_run("_" * 95)
    rr.font.size = Pt(6)
    rr.font.color.rgb = ACCENT


def build_resume(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Pt(40)
        sec.left_margin = sec.right_margin = Pt(54)

    _name_block(doc, DATA["currentTitle"])

    _section(doc, "Professional Summary")
    doc.add_paragraph(DATA["summary"])

    _section(doc, "Technical Skills")
    doc.add_paragraph(" • ".join(DATA["skills"]))

    _section(doc, "Professional Experience")
    for job in DATA["workHistory"]:
        head = doc.add_paragraph()
        h = head.add_run(f"{job['title']} — {job['company']}")
        h.bold = True
        h.font.size = Pt(11)
        meta = head.add_run(f"   ({job['startDate']} – {job['endDate']}, {job['location']})")
        meta.italic = True
        meta.font.size = Pt(9)
        for hl in job["highlights"]:
            b = doc.add_paragraph(hl, style="List Bullet")
            b.paragraph_format.space_after = Pt(1)

    _section(doc, "Education")
    for ed in DATA["education"]:
        p = doc.add_paragraph()
        r = p.add_run(f"{ed['degree']} — {ed['school']}")
        r.bold = True
        p.add_run(f"   ({ed['graduationDate']}, {ed['location']})").italic = True
        if ed.get("details"):
            doc.add_paragraph(ed["details"])

    _section(doc, "Certifications")
    for cert in DATA["certifications"]:
        doc.add_paragraph(cert, style="List Bullet")

    doc.save(path)


def build_cover_letter(path):
    posting = DATA["_meta"]["tailoredToPosting"]
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Pt(54)
        sec.left_margin = sec.right_margin = Pt(72)

    _name_block(doc, "Senior .NET Engineer — Application")
    doc.add_paragraph()

    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    doc.add_paragraph(f"Hiring Team\n{posting['company']}\n{posting['location']}")
    doc.add_paragraph()

    doc.add_paragraph("Dear FlashTec Hiring Team,")

    body = [
        (
            f"I am excited to apply for the {posting['title']} position "
            f"({posting['jobId']}) on FlashTec's Platform team in {posting['location']}. "
            f"{DATA['whyInterested']}"
        ),
        (
            f"Over the past {DATA['totalExperienceYears']} years I have built high-throughput, "
            "low-latency payment and ledger systems in C# and .NET. In my current role as "
            f"{DATA['currentTitle']} at {DATA['currentCompany']}, I design and operate .NET 8 "
            "microservices that process over three million transactions a day at sub-100ms p99 "
            "latency, model complex financial domains with Domain-Driven Design, and run "
            "event-driven pipelines on Azure Service Bus and Kafka. This maps directly to your "
            "need for production ASP.NET Core Web APIs, strong SQL Server skills, and "
            "message-based architecture experience."
        ),
        (
            "Beyond the code, I own features end-to-end — architecture through on-call — and "
            "mentor mid-level engineers, having introduced clean-architecture and "
            "automated-testing standards across my team. FlashTec's five-nines, "
            "Microsoft-centric stack is precisely where I do my best work."
        ),
        (
            f"I am authorized to work in the U.S. ({DATA['workAuthorization']}), based in "
            f"{DATA['city']}, {DATA['state']}, and available to start "
            f"{DATA['availableStartDate']}. I would welcome the chance to discuss how I can "
            "help FlashTec keep money moving reliably at scale. Thank you for your "
            "consideration."
        ),
    ]
    for para in body:
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    sig = doc.add_paragraph()
    sr = sig.add_run(DATA["fullName"])
    sr.bold = True

    doc.save(path)


if __name__ == "__main__":
    rp = os.path.join(HERE, "jason_statham_resume.docx")
    cp = os.path.join(HERE, "jason_statham_cover_letter.docx")
    build_resume(rp)
    build_cover_letter(cp)
    print("Wrote:", rp)
    print("Wrote:", cp)
